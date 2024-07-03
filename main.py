import telebot
import os
import nltk
import docx
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton, ReplyKeyboardMarkup, KeyboardButton

# Replace YOUR_BOT_TOKEN with your actual token
BOT_TOKEN = '7308978410:AAGknd4DRytR_I94f3yJHqAPtn7Tv034EwU'

# Download NLTK data files (run this once)
nltk.download('punkt')

# Create a Bot instance
bot = telebot.TeleBot(BOT_TOKEN)

def read_docx_file(filename):
    try:
        doc = docx.Document(filename)
        title = doc.paragraphs[0].text.strip() if doc.paragraphs else "عنوان غير موجود"
        full_text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        return title, full_text
    except Exception as e:
        print(f"Error reading {filename}: {e}")
        return "خطأ في قراءة الملف", ""

def extract_articles(text, search_term):
    articles = text.split('\n\n')
    matching_articles = [article.strip() for article in articles if search_term in article]
    return matching_articles

def send_large_message(chat_id, message):
    message_chunks = [message[i:i+4096] for i in range(0, len(message), 4096)]
    for chunk in message_chunks:
        bot.send_message(chat_id, chunk)

def generate_keyboard():
    folder_path = 'Q/Docx'
    keyboard_markup = ReplyKeyboardMarkup(resize_keyboard=True)

    search_all_button = KeyboardButton("بحث شامل")
    keyboard_markup.add(search_all_button)

    for idx, filename in enumerate(os.listdir(folder_path), start=1):
        if filename.endswith('.docx'):
            button_text = f"{idx}. {os.path.splitext(filename)[0].strip()}"
            button = KeyboardButton(text=button_text)
            keyboard_markup.add(button)

    return keyboard_markup

@bot.message_handler(commands=['start'])
def start_message(message):
    username = message.from_user.first_name
    user_id = message.from_user.id
    phone_number = message.contact.phone_number if message.contact else "غير متوفر"
    user_username = f"@{message.from_user.username}" if message.from_user.username else "غير متوفر"
    
    bot.send_message(message.chat.id, f"مرحباً {username}! أهلاً بك.")
    bot.send_message(message.chat.id, "*يرجى اختيار البحث المخصص أو اختيار 'بحث شامل' من الأزرار:*", reply_markup=generate_keyboard())
    
    # إرسال التفاصيل إلى حساب المطور
    developer_chat_id = 5063004694  # استبدل هذا الرقم بمعرف المطور
    bot.send_message(developer_chat_id, 
        f"تم تسجيل الدخول من قبل :-\n\n"
        f"اسم المستخدم / {username}\n"
        
        f"معرف المستخدم / {user_username}\n"
        
        f"رقم المستخدم / {phone_number}\n"
        
        f"كود المستخدم / {user_id}"
    )

@bot.message_handler(func=lambda message: True)
def handle_message(message):
    folder_path = 'Q/Docx'
    valid_filenames = [os.path.splitext(f)[0].strip() for f in os.listdir(folder_path) if f.endswith('.docx')]
    if message.text == "بحث شامل":
        bot.send_message(message.chat.id, "يرجى إدخال جزء أو نص المادة المراد البحث عنها:")
        bot.register_next_step_handler(message, search_all_documents)
    elif any(message.text.startswith(f"{idx}. {name}") for idx, name in enumerate(valid_filenames, start=1)):
        selected_filename = message.text.split(". ", 1)[1] + ".docx"
        bot.send_message(message.chat.id, "يرجى إدخال جزء أو نص المادة المراد البحث عنها في المستند المحدد:")
        bot.register_next_step_handler(message, lambda msg: search_in_specific_document(msg, selected_filename))
    else:
        markup = InlineKeyboardMarkup()
        home_button = InlineKeyboardButton("العودة للقائمة الرئيسة", callback_data="go_home")
        markup.add(home_button)
        bot.send_message(message.chat.id, "الرجاء اختيار مستند صالح أو 'بحث شامل'.", reply_markup=markup)

def search_all_documents(message):
    search_term = message.text
    folder_path = 'Q/Docx'
    found_any = False
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            filepath = os.path.join(folder_path, filename)
            title, full_text = read_docx_file(filepath)
            found_articles = extract_articles(full_text, search_term)
            if found_articles:
                found_any = True
                responses = []
                for article in found_articles:
                    file_base_name = os.path.splitext(filename)[0].strip()
                    formatted_response = f"⭕️ - {title}\n بشأن : {file_base_name}\n\n{article}"
                    responses.append(formatted_response)
                response = "\n\n".join(responses)
                send_large_message(message.chat.id, response)
                markup = InlineKeyboardMarkup()
                send_button = InlineKeyboardButton("أرسل المستند", callback_data=f"send_doc:{filename}")
                home_button = InlineKeyboardButton("العودة للقائمة الرئيسة", callback_data="go_home")
                markup.add(send_button, home_button)
                bot.send_message(message.chat.id, "هل ترغب في استلام المصدر بالكامل؟", reply_markup=markup)
    if not found_any:
        markup = InlineKeyboardMarkup()
        home_button = InlineKeyboardButton("العودة للقائمة الرئيسة", callback_data="go_home")
        markup.add(home_button)
        bot.send_message(message.chat.id, f"لم أجد أي مستند يحتوي على '{search_term}' في المجلد.", reply_markup=markup)

def search_in_specific_document(message, filename):
    search_term = message.text
    filepath = os.path.join('Q/Docx', filename)
    title, full_text = read_docx_file(filepath)
    found_articles = extract_articles(full_text, search_term)
    if found_articles:
        responses = []
        for article in found_articles:
            formatted_response = f"⭕️ - {title}\n\n{article}"
            responses.append(formatted_response)
        response = "\n\n".join(responses)
        send_large_message(message.chat.id, response)
        markup = InlineKeyboardMarkup()
        send_button = InlineKeyboardButton("أرسل المستند", callback_data=f"send_doc:{filename}")
        home_button = InlineKeyboardButton("العودة للقائمة الرئيسة", callback_data="go_home")
        markup.add(send_button, home_button)
        bot.send_message(message.chat.id, "هل ترغب في استلام المصدر بالكامل؟", reply_markup=markup)
    else:
        markup = InlineKeyboardMarkup()
        home_button = InlineKeyboardButton("العودة للقائمة الرئيسة", callback_data="go_home")
        markup.add(home_button)
        bot.send_message(message.chat.id, f"لم أجد أي مستند يحتوي على '{search_term}'.", reply_markup=markup)

@bot.callback_query_handler(func=lambda call: call.data.startswith("send_doc:"))
def callback_send_doc(call):
    filename = call.data.split(":")[1]
    filepath = os.path.join('Q/Docx', filename)
    if os.path.exists(filepath):
        with open(filepath, 'rb') as doc_file:
            bot.send_document(call.message.chat.id, doc_file)
    else:
        bot.send_message(call.message.chat.id, "عذراً، لم أتمكن من العثور على المستند.")

@bot.callback_query_handler(func=lambda call: call.data == "go_home")
def callback_go_home(call):
    start_message(call.message)

# Start the bot
bot.polling(none_stop=True)
